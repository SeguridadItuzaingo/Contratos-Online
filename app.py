from drive_util import upload_path_to_drive
from flask import Flask, render_template, request, send_file, session
from uuid import uuid4
import os
import re
import unicodedata
from datetime import datetime
from io import BytesIO
import base64
import html

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from dotenv import load_dotenv

# Email helper (debe existir correo_util.py con enviar_email(to, asunto, cuerpo, adjunto_path))
import correo_util

# (Opcional/Windows) Para docx2pdf con Word
try:
    import pythoncom
    HAS_PYTHONCOM = True
except Exception:
    HAS_PYTHONCOM = False

# =========================================================
# App / Config
# =========================================================
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET", os.urandom(24))

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")
TEMPLATE_DOCX = os.path.join(BASE_DIR, "Contrato_Plantilla.docx")

# Asegurar carpeta de salida también en Gunicorn (Render)
os.makedirs(STATIC_DIR, exist_ok=True)

# Datos empresa (ajustables)
RESPONSABLE_EMPRESA = "Alan Arndt, Dueño de la Empresa"
FIRMA_EMPRESA_PATH = os.path.join(STATIC_DIR, "firma_empresa.png")
EMAIL_EMPRESA = os.getenv("EMAIL_EMPRESA", "")
CONTACTO_TELEFONO = os.getenv("CONTACTO_TELEFONO", "")

# Tamaños de firma
SIGNATURE_IMAGE_WIDTH_IN = 2.8
SIGNATURE_LABEL_FONT_PT = 12

# =========================================================
# Helpers
# =========================================================
def _now_tag() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def _slug(s: str) -> str:
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[^a-zA-Z0-9]+", "_", s).strip("_").lower()
    return s or f"cliente_{uuid4().hex[:6]}"

def _docx_to_html_paragraphs(path_docx: str) -> str:
    """Convierte el DOCX a párrafos HTML simples para previsualizarlo."""
    doc = Document(path_docx)
    parts = []
    for p in doc.paragraphs:
        text = html.escape(p.text).strip()
        parts.append(f"<p>{text or '&nbsp;'}</p>")
    return "\n".join(parts)

def _b64_to_pil_image(b64data: str) -> Image.Image:
    """Convierte base64 (data:image/png;base64,...) a PIL Image RGBA."""
    if "," in b64data:
        b64data = b64data.split(",", 1)[1]
    raw = base64.b64decode(b64data)
    return Image.open(BytesIO(raw)).convert("RGBA")

def _insert_text_placeholders(doc: Document, mapping: dict):
    """
    Reemplazo robusto a nivel de párrafo/celda.
    - Normaliza NBSP/zero-width en el texto del DOCX.
    - Hace reemplazos literales (mapping).
    - Aplica un regex tolerante para {{ ubicacion_monitoreo }} por si quedó cortado/espaciado.
    """
    import re

    def _norm(s: str) -> str:
        # NBSP -> espacio, elimina zero-width
        return (s or "").replace("\xa0", " ").replace("\u200b", "")

    # regex tolerante: {{ u b i c a c i o n _? m o n i t o r e o }}
    ubi_mon_re = re.compile(
        r"\{\{\s*u\s*b\s*i\s*c\s*a\s*c\s*i\s*o\s*n\s*(?:_| |\t)?\s*m\s*o\s*n\s*i\s*t\s*o\s*r\s*e\s*o\s*\}\}",
        re.IGNORECASE
    )

    def replace_text(text: str) -> str:
        t = _norm(text)
        # Pase literal con lo que ya tenés en 'mapping'
        for k, v in mapping.items():
            t = t.replace(k, v)
        # Pase extra solo para ubicacion_monitoreo tolerando cortes/espacios
        valor_ubi_mon = mapping.get("{{ ubicacion_monitoreo }}", "")
        if valor_ubi_mon:
            t = ubi_mon_re.sub(valor_ubi_mon, t)
        return t

    # Párrafos fuera de tablas
    for p in doc.paragraphs:
        new = replace_text(p.text)
        if new != p.text:
            p.text = new

    # Celdas de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new = replace_text(p.text)
                    if new != p.text:
                        p.text = new

def _ensure_company_signature(path_png: str, texto: str = "Alan Arndt"):
    """Genera una firma PNG básica para la empresa si no existe."""
    if os.path.exists(path_png):
        return
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGBA", (600, 220), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 72)
    except Exception:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), texto, font=font)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((600 - w) // 2, (220 - h) // 2), texto, fill=(0, 0, 0, 255), font=font)
    img.save(path_png, "PNG")

def _add_signatures_section(doc: Document, firma_cliente_path: str, firma_empresa_path: str):
    """
    Inserta ambas firmas en una tabla 2x2:
      Fila 0: imágenes (cliente | empresa)
      Fila 1: rótulos  ("Firma del Cliente" | "Firma de la Empresa")
    Si detecta ya una tabla con esos rótulos, la reutiliza.
    """
    target = None
    for tbl in doc.tables:
        try:
            if len(tbl.rows) >= 2 and len(tbl.columns) >= 2:
                if ("firma del cliente" in tbl.cell(1, 0).text.lower()
                        and "firma de la empresa" in tbl.cell(1, 1).text.lower()):
                    target = tbl
                    break
        except Exception:
            continue

    if target is None:
        target = doc.add_table(rows=2, cols=2)
        target.cell(1, 0).text = "Firma del Cliente"
        target.cell(1, 1).text = "Firma de la Empresa"

    while len(target.rows) < 2:
        target.add_row()

    # Layout fijo
    try:
        target.autofit = False
    except Exception:
        pass
    try:
        tbl = target._tbl
        tblPr = tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
    except Exception:
        pass

    # Ancho parejo
    col_w = Inches(3.2)
    for i in range(2):
        try:
            target.columns[i].width = col_w
        except Exception:
            pass
        for cell in target.columns[i].cells:
            try:
                cell.width = col_w
            except Exception:
                pass

    # Fila 0: imágenes centradas
    for col, img_path in enumerate([firma_cliente_path, firma_empresa_path]):
        c = target.cell(0, col)
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(img_path) and os.path.getsize(img_path) > 0:
            p.add_run().add_picture(img_path, width=Inches(SIGNATURE_IMAGE_WIDTH_IN))

    # Fila 1: rótulos centrados
    for col, texto in enumerate(["Firma del Cliente", "Firma de la Empresa"]):
        c = target.cell(1, col)
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.bold = True
        try:
            run.font.size = Pt(SIGNATURE_LABEL_FONT_PT)
        except Exception:
            pass

def _export_to_pdf_safe(out_docx: str, out_pdf: str) -> bool:
    """
    Convierte DOCX -> PDF en Linux usando LibreOffice.
    - Fuerza headless/locale y prueba 'soffice' y 'libreoffice'.
    """
    import subprocess, shlex

    out_dir = os.path.dirname(out_pdf) or "."
    env = os.environ.copy()
    env.setdefault("HOME", "/tmp")
    env.setdefault("LANG", "en_US.UTF-8")

    cmds = [
        f"soffice --headless --norestore --nolockcheck --nodefault "
        f"--convert-to pdf:writer_pdf_Export --outdir {shlex.quote(out_dir)} {shlex.quote(out_docx)}",
        f"libreoffice --headless --norestore --nolockcheck --nodefault "
        f"--convert-to pdf:writer_pdf_Export --outdir {shlex.quote(out_dir)} {shlex.quote(out_docx)}",
    ]

    for cmd in cmds:
        try:
            subprocess.run(
                cmd, shell=True, check=True,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                timeout=180, env=env
            )
            gen_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(out_docx))[0] + ".pdf")
            if os.path.exists(gen_pdf) and os.path.getsize(gen_pdf) > 0:
                if gen_pdf != out_pdf:
                    os.replace(gen_pdf, out_pdf)
                return True
        except Exception:
            continue

    return False

# =========================================================
# Rutas
# =========================================================
@app.route("/", methods=["GET"])
def index():
    if not os.path.exists(TEMPLATE_DOCX):
        contrato_html = "<p><strong>No se encontró la plantilla del contrato.</strong></p>"
    else:
        contrato_html = _docx_to_html_paragraphs(TEMPLATE_DOCX)
    return render_template("formulario_contrato.html", contrato_html=contrato_html)

@app.route("/generar", methods=["POST"])
def generar():
    # 1) Validación básica
nombre = request.form.get("nombre", "").strip()
dni = request.form.get("dni", "").strip()
email = (request.form.get("email") or request.form.get("correo") or request.form.get("mail") or "").strip()

# Domicilio del abonado
ubicacion = request.form.get("ubicacion", "").strip()

# Lugar monitoreado
ubicacion_monitoreo = request.form.get("ubicacion_monitoreo", "").strip()

firma_b64 = (request.form.get("firmaBase64") or request.form.get("firma") or "").strip()

# (Opcional) Log rápido para verificar que llegan los dos valores
app.logger.info(f"[DBG] ubicacion='{ubicacion}' | ubicacion_monitoreo='{ubicacion_monitoreo}'")

    faltantes = [k for k, v in {
    "nombre": nombre,
    "dni": dni,
    "email": email,
    "ubicacion": ubicacion,                       # domicilio del abonado
    "ubicacion_monitoreo": ubicacion_monitoreo,   # lugar monitoreado
    "firma": firma_b64
}.items() if not v]
if faltantes:
    return f"Faltan campos obligatorios: {', '.join(faltantes)}", 400
    
    # 2) Preparar rutas de salida
    slug = f"{_slug(nombre)}_{_now_tag()}_{uuid4().hex[:6]}"
    out_docx = os.path.join(STATIC_DIR, f"{slug}.docx")
    out_pdf = os.path.join(STATIC_DIR, f"{slug}.pdf")
    firma_cliente_path = os.path.join(STATIC_DIR, f"firma_{slug}.png")

    # 3) Guardar firma del cliente
    try:
        img = _b64_to_pil_image(firma_b64)
        img.save(firma_cliente_path, "PNG")
    except Exception:
        return "La imagen de la firma es inválida.", 400

    # 3.b) Abrir plantilla y hacer reemplazos
    try:
        doc = Document(TEMPLATE_DOCX)
    except Exception as e:
        return f"No se pudo abrir la plantilla del contrato: {e}", 500

    mapping = {
    "{{ nombre }}": nombre,
    "{{ dni }}": dni,
    "{{ email }}": email,
    "{{ ubicacion }}": ubicacion,                         # domicilio del abonado
    "{{ ubicacion_monitoreo }}": ubicacion_monitoreo,     # lugar monitoreado
    "{{ fecha_hoy }}": datetime.now().strftime("%d/%m/%Y"),
}
_insert_text_placeholders(doc, mapping)
    
    # 3.c) Firmas (cliente + empresa)
    _ensure_company_signature(FIRMA_EMPRESA_PATH)
    _add_signatures_section(doc, firma_cliente_path, FIRMA_EMPRESA_PATH)

    # 4) Guardar DOCX
    try:
        doc.save(out_docx)
    except Exception as e:
        return f"No se pudo guardar el DOCX: {e}", 500

    # 5) Convertir a PDF con fallback
    pdf_ok = _export_to_pdf_safe(out_docx, out_pdf)
    if pdf_ok and os.path.exists(out_pdf) and os.path.getsize(out_pdf) > 0:
        session["archivo_pdf"] = out_pdf
    else:
        session["archivo_pdf"] = out_docx  # fallback sin romper el flujo

    # 5.1) Subir a Google Drive (no bloquea al usuario si falla)
    try:
        adjunto_path = session["archivo_pdf"]
        ext = os.path.splitext(adjunto_path)[1].lower()
        mimetype = (
            "application/pdf"
            if ext == ".pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Nombre prolijo para el archivo remoto
        safe_nombre = "".join(c for c in nombre if c.isalnum() or c in " _-").strip().replace(" ", "_")
        fecha_str = datetime.now().strftime("%Y-%m-%d")
        nombre_remoto = f"{fecha_str}_{safe_nombre}_{dni}_Contrato{ext}"

        drive_id = upload_path_to_drive(adjunto_path, nombre_remoto, mimetype)
        app.logger.info(f"[Drive] Subido OK. fileId={drive_id}")
    except Exception as e:
        app.logger.exception(f"[Drive] Falló la subida: {e}")

    # 6) Enviar correos (cliente y empresa) sin romper flujo si falla SMTP
    try:
        adjunto = session["archivo_pdf"]
        asunto = "Contrato firmado - Seguridad Ituzaingó"
cuerpo = (
    f"Estimado/a {nombre},\n\n"
    f"Adjuntamos el contrato firmado correspondiente al servicio de monitoreo en {ubicacion_monitoreo}.\n"
    f"Domicilio del abonado: {ubicacion}.\n"
    "Le recomendamos conservar el archivo para su referencia.\n\n"
    "Quedamos a disposición por cualquier consulta.\n\n"
    "Atentamente,\n"
    "Seguridad Ituzaingó\n"
    "Alan Arndt — Dueño de la Empresa\n"
    f"Tel.: {CONTACTO_TELEFONO or '-'}\n"
    f"Email: {EMAIL_EMPRESA or '-'}\n"
)

        # Enviar al cliente
        if email:
            correo_util.enviar_email(email, asunto, cuerpo, adjunto)

        # Copia a la empresa
        if EMAIL_EMPRESA:
            correo_util.enviar_email(
                EMAIL_EMPRESA,
                "Nuevo contrato firmado - Seguridad Ituzaingó",
                f"El/La cliente {nombre} firmó un contrato.\n\n{cuerpo}",
                adjunto
            )
    except Exception:
        # No interrumpir la UX si falla el SMTP
        pass

    # 7) Página de agradecimiento con descarga
    return render_template("agradecimiento.html", telefono=CONTACTO_TELEFONO)

@app.route("/descargar", methods=["GET"])
def descargar():
    archivo = session.get("archivo_pdf")
    if not archivo or not os.path.exists(archivo):
        return "Error: No se encontró el contrato para descargar.", 404
    return send_file(archivo, as_attachment=True, download_name=os.path.basename(archivo))

# =========================================================
# Main (solo desarrollo; en Render se usa gunicorn)
# =========================================================
if __name__ == "__main__":
    app.run(debug=True)










